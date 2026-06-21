#!/usr/bin/env python3
"""
Universal Markdown Converters
Five format converters with shared utilities for consistent output.
"""

import hashlib
import random
import re
import shutil
import time
from datetime import date, datetime, timezone
from email.utils import parsedate_to_datetime
from pathlib import Path
from typing import NamedTuple
from urllib.parse import urlsplit

import os
import ssl

import certifi

# Fix SSL certificate discovery on macOS + Python 3.14 + OpenSSL 3.6:
# certifi's bundled CA file may be rejected by newer OpenSSL builds.
# Prefer the system/Homebrew CA file that OpenSSL ships with; fall back
# to certifi only when no system file is found.
def _find_ca_file() -> str:
    """Return the best available CA certificate bundle path."""
    # 1. Already set by user/environment
    env_ca = os.environ.get("SSL_CERT_FILE")
    if env_ca and os.path.isfile(env_ca):
        return env_ca
    # 2. System OpenSSL default (works with Homebrew OpenSSL 3.x)
    _paths = ssl.get_default_verify_paths()
    for candidate in (_paths.cafile, _paths.openssl_cafile):
        if candidate and os.path.isfile(candidate):
            return candidate
    # 3. Fallback to certifi
    return certifi.where()

_CA_FILE = _find_ca_file()
os.environ["SSL_CERT_FILE"] = _CA_FILE

import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.table import Table as DocxTable
from markdownify import markdownify as html_to_md
from spreadsheet_converter import write_xlsx_sheets
from striprtf.striprtf import rtf_to_text


WEB_REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/136.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

WEB_FETCH_MAX_ATTEMPTS = 3
WEB_FETCH_BASE_DELAY_SECONDS = 1.0
WEB_FETCH_MAX_DELAY_SECONDS = 8.0
WEB_FETCH_JITTER_SECONDS = 0.25
WEB_FETCH_RETRYABLE_STATUSES = {429, 500, 502, 503, 504}
WEB_FETCH_RETRYABLE_EXCEPTIONS = (requests.ConnectionError, requests.Timeout)


def _parse_retry_after(value: str | None) -> float | None:
    """Parse Retry-After seconds or HTTP date into a sleep delay."""
    if not value:
        return None

    value = value.strip()
    try:
        seconds = float(value)
    except ValueError:
        try:
            retry_at = parsedate_to_datetime(value)
        except (TypeError, ValueError, IndexError, OverflowError):
            return None
        if retry_at.tzinfo is None:
            retry_at = retry_at.replace(tzinfo=timezone.utc)
        seconds = (retry_at - datetime.now(timezone.utc)).total_seconds()

    return max(seconds, 0.0)


def _retry_delay_seconds(attempt: int) -> float:
    """Calculate a short exponential backoff with a small amount of jitter."""
    base_delay = min(WEB_FETCH_BASE_DELAY_SECONDS * (2 ** (attempt - 1)), WEB_FETCH_MAX_DELAY_SECONDS)
    return base_delay + random.uniform(0.0, WEB_FETCH_JITTER_SECONDS)


def _fetch_url_html(url: str) -> str:
    """Fetch HTML with conservative retries for transient failures only."""
    for attempt in range(1, WEB_FETCH_MAX_ATTEMPTS + 1):
        try:
            resp = requests.get(url, timeout=30, headers=WEB_REQUEST_HEADERS, verify=_CA_FILE)
        except WEB_FETCH_RETRYABLE_EXCEPTIONS:
            if attempt == WEB_FETCH_MAX_ATTEMPTS:
                raise
            time.sleep(_retry_delay_seconds(attempt))
            continue

        if resp.status_code in WEB_FETCH_RETRYABLE_STATUSES and attempt < WEB_FETCH_MAX_ATTEMPTS:
            retry_after = _parse_retry_after(resp.headers.get("Retry-After"))
            time.sleep(retry_after if retry_after is not None else _retry_delay_seconds(attempt))
            continue

        resp.raise_for_status()
        return resp.text

    raise RuntimeError("URL fetch retry loop exited unexpectedly")


# ---------------------------------------------------------------------------
# Shared types
# ---------------------------------------------------------------------------

class ConvertResult(NamedTuple):
    success: bool
    output_path: str
    word_count: int
    message: str


# ---------------------------------------------------------------------------
# Shared utilities (extracted from pdf_to_md / docx_to_md)
# ---------------------------------------------------------------------------

def safe_filename(name: str) -> str:
    """Sanitize a document title into a lowercase-dash filename."""
    safe = re.sub(r'[^\w\s\-]', '', name).strip()
    return re.sub(r'\s+', '-', safe).lower()


def output_stem(title: str, source_file: str, source_type: str) -> str:
    """Build a stable output stem, avoiding collisions for URL-based pages."""
    stem = safe_filename(title) or "converted"

    if source_type != "html" or not source_file.startswith(("http://", "https://")):
        return stem

    parsed = urlsplit(source_file)
    tail = Path(parsed.path).name.replace(".", " ")
    tail_slug = safe_filename(tail)
    url_hash = hashlib.sha1(source_file.encode("utf-8")).hexdigest()[:8]
    suffix = f"{tail_slug}-{url_hash}" if tail_slug else url_hash
    return f"{stem}-{suffix}"


def normalize_blanks(text: str) -> str:
    """Collapse 3+ consecutive newlines into 2."""
    return re.sub(r'\n{3,}', '\n\n', text)


def build_header(title: str, source: str, word_count: int, **extras) -> str:
    """Build the blockquote metadata header used by all converters."""
    lines = [f"# {title}", ""]
    lines.append(f"> **Source**: `{source}`  ")
    for key, val in extras.items():
        lines.append(f"> **{key}**: {val}  ")
    lines.append(f"> **Word Count**: {word_count:,}  ")
    lines.extend(["", "---", ""])
    return "\n".join(lines)


def vault_frontmatter(title: str, source_type: str, source_file: str) -> str:
    """Generate Obsidian-compatible YAML frontmatter."""
    return (
        f"---\n"
        f'title: "{title}"\n'
        f"tags: [converted, {source_type}]\n"
        f"created: {date.today().isoformat()}\n"
        f'source: "{source_file}"\n'
        f"---\n\n"
    )


def write_output(
    body: str,
    title: str,
    source_file: str,
    word_count: int,
    output_dir: Path,
    source_type: str,
    vault_dir: Path | None = None,
    header_extras: dict | None = None,
) -> ConvertResult:
    """Write markdown to output_dir and optionally copy to vault."""
    output_dir.mkdir(parents=True, exist_ok=True)
    md_name = f"{output_stem(title, source_file, source_type)}.md"
    md_path = output_dir / md_name

    header = build_header(title, source_file, word_count, **(header_extras or {}))
    content = header + body + "\n"
    md_path.write_text(content, encoding="utf-8")

    # Vault delivery
    if vault_dir:
        vault_type_dir = vault_dir / source_type
        vault_type_dir.mkdir(parents=True, exist_ok=True)
        vault_path = vault_type_dir / md_name
        vault_content = vault_frontmatter(title, source_type, source_file) + content
        vault_path.write_text(vault_content, encoding="utf-8")

    return ConvertResult(True, str(md_path), word_count, f"OK -> {md_name}")


# ---------------------------------------------------------------------------
# Format routing
# ---------------------------------------------------------------------------

SUPPORTED = {'.pdf', '.docx', '.html', '.htm', '.txt', '.rtf', '.xlsx'}

SUBFOLDER = {
    '.pdf': 'pdf', '.docx': 'docx',
    '.html': 'html', '.htm': 'html',
    '.txt': 'txt', '.rtf': 'rtf',
    '.xlsx': 'spreadsheets',
}


def route(path: str, base_output: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Detect format and call the right converter."""
    # URL detection
    if path.startswith("http://") or path.startswith("https://"):
        out = base_output / "html"
        return convert_html(path, out, vault_dir)

    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED:
        return ConvertResult(False, "", 0, f"Unsupported format: {ext}")

    out = base_output / SUBFOLDER[ext]
    converters = {
        '.pdf': convert_pdf,
        '.docx': convert_docx,
        '.html': convert_html,
        '.htm': convert_html,
        '.txt': convert_txt,
        '.rtf': convert_rtf,
        '.xlsx': convert_xlsx,
    }
    return converters[ext](path, out, vault_dir)


# ---------------------------------------------------------------------------
# 1. PDF converter (with OCR auto-fallback)
# ---------------------------------------------------------------------------

def convert_pdf(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert a PDF to Markdown. Falls back to OCR if no selectable text."""
    p = Path(path)
    doc = fitz.open(path)
    pages = []
    total_words = 0
    page_count = len(doc)

    for i in range(page_count):
        text = doc[i].get_text("text")
        if text.strip():
            pages.append((i + 1, text))
            total_words += len(text.split())
    doc.close()

    # Auto-fallback to OCR
    if total_words == 0:
        return _convert_pdf_ocr(path, output_dir, vault_dir)

    body_lines = []
    for num, text in pages:
        body_lines.append(f"## Page {num}\n")
        body_lines.append(normalize_blanks(text.strip()))
        body_lines.append("")

    body = "\n".join(body_lines)
    return write_output(
        body, p.stem, p.name, total_words, output_dir, "pdf", vault_dir,
        header_extras={"Pages": str(page_count)},
    )


def _convert_pdf_ocr(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """OCR fallback for scanned PDFs."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ConvertResult(False, "", 0, "ERROR: pytesseract/Pillow not installed for OCR")

    p = Path(path)
    doc = fitz.open(path)
    page_count = len(doc)
    pages = []
    total_words = 0
    zoom = 300 / 72
    matrix = fitz.Matrix(zoom, zoom)
    start = time.time()

    for i in range(page_count):
        pix = doc[i].get_pixmap(matrix=matrix)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        if text.strip():
            pages.append((i + 1, text))
            total_words += len(text.split())
    doc.close()
    elapsed = time.time() - start

    if total_words == 0:
        return ConvertResult(False, "", 0, "ERROR: No text extracted even with OCR")

    body_lines = []
    for num, text in pages:
        body_lines.append(f"## Page {num}\n")
        body_lines.append(normalize_blanks(text.strip()))
        body_lines.append("")

    body = "\n".join(body_lines)
    return write_output(
        body, p.stem, p.name, total_words, output_dir, "pdf", vault_dir,
        header_extras={
            "Pages": str(page_count),
            "Extracted via": "Tesseract OCR (local)",
            "Processing time": f"{elapsed:.1f}s",
        },
    )


# ---------------------------------------------------------------------------
# 2. DOCX converter (ported from docx_to_md.py)
# ---------------------------------------------------------------------------

def _run_to_md(run) -> str:
    text = run.text
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def _para_to_md(para) -> str:
    style = para.style.name.lower()
    parts = [_run_to_md(r) for r in para.runs]
    text = "".join(parts).strip() or para.text.strip()
    if not text:
        return ""

    if style.startswith("heading"):
        try:
            level = min(int(style.split()[-1]), 6)
        except (ValueError, IndexError):
            level = 1
        return f"{'#' * level} {text}"
    if style == "title":
        return f"# {text}"
    if style == "subtitle":
        return f"## {text}"
    if style.startswith("list bullet"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}- {text}"
    if style.startswith("list number"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}1. {text}"
    if "quote" in style:
        return f"> {text}"
    return text


def _table_to_md(table: DocxTable) -> str:
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * cols) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_docx(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert a DOCX file preserving headings, formatting, lists, and tables."""
    p = Path(path)
    doc = Document(path)
    blocks = []
    word_count = 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            for para in doc.paragraphs:
                if para._element is child:
                    line = _para_to_md(para)
                    blocks.append(line if line else "")
                    if line:
                        word_count += len(line.split())
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._tbl is child:
                    md = _table_to_md(table)
                    if md:
                        blocks.extend(["", md, ""])
                        word_count += len(md.split())
                    break

    body = normalize_blanks("\n".join(blocks)).strip()
    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(body, p.stem, p.name, word_count, output_dir, "docx", vault_dir)


# ---------------------------------------------------------------------------
# 3. HTML / URL converter
# ---------------------------------------------------------------------------

def convert_html(path_or_url: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert HTML file or URL to Markdown."""
    is_url = path_or_url.startswith("http://") or path_or_url.startswith("https://")

    if is_url:
        html = _fetch_url_html(path_or_url)
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string.strip() if soup.title and soup.title.string else path_or_url
        source = path_or_url
    else:
        p = Path(path_or_url)
        html = p.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string.strip() if soup.title and soup.title.string else p.stem
        source = p.name

    # Remove script/style tags before conversion
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    body = html_to_md(str(soup), heading_style="ATX", strip=["img"])
    body = normalize_blanks(body).strip()
    word_count = len(body.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty page)")

    extras = {"URL": f"`{path_or_url}`"} if is_url else {}
    return write_output(body, title, source, word_count, output_dir, "html", vault_dir, extras)


# ---------------------------------------------------------------------------
# 4. TXT converter
# ---------------------------------------------------------------------------

def convert_txt(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Wrap a plain text file in a Markdown metadata header."""
    p = Path(path)
    text = p.read_text(encoding="utf-8", errors="replace").strip()
    word_count = len(text.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(text, p.stem, p.name, word_count, output_dir, "txt", vault_dir)


# ---------------------------------------------------------------------------
# 5. RTF converter
# ---------------------------------------------------------------------------

def convert_rtf(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert RTF to Markdown by stripping RTF formatting."""
    p = Path(path)
    try:
        # Read as raw bytes first — RTF files are almost never UTF-8.
        raw_bytes = p.read_bytes()

        # Decode with latin-1, which maps every byte 0x00-0xFF one-to-one to
        # Unicode code-points. This preserves the original bytes so that
        # striprtf can interpret RTF encoding directives (\\ansicpg, etc.)
        # and produce correct Unicode output.
        raw = raw_bytes.decode("latin-1")

        text = rtf_to_text(raw, errors="ignore").strip()
    except Exception as e:
        return ConvertResult(False, "", 0, f"ERROR reading RTF: {e}")

    word_count = len(text.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(text, p.stem, p.name, word_count, output_dir, "rtf", vault_dir)


def convert_xlsx(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert each XLSX workbook sheet to a separate Markdown file."""
    sheet_count, total_words, output_paths = write_xlsx_sheets(
        path,
        output_dir,
        vault_dir,
        write_output,
    )

    if sheet_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty workbook)")

    return ConvertResult(
        True,
        str(output_dir),
        total_words,
        f"OK -> {sheet_count} sheets: {', '.join(output_paths)}",
    )


# ---------------------------------------------------------------------------
# 7. Raw pasted text converter
# ---------------------------------------------------------------------------

def convert_raw_text(text: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Save raw pasted text as a Markdown file."""
    text = text.strip()
    word_count = len(text.split())
    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    # Generate title from first line (truncated to 60 chars)
    first_line = text.split('\n')[0].strip()
    title = first_line[:60] if first_line else "Pasted Text"
    # Clean title for display
    title = re.sub(r'[#*>\-=]', '', title).strip() or "Pasted Text"

    return write_output(text, title, "pasted-text", word_count, output_dir, "txt", vault_dir)


# ---------------------------------------------------------------------------
# 8. Pasted input router (text or URL)
# ---------------------------------------------------------------------------

def convert_pasted(text: str, base_output: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Route pasted text - if it looks like a URL, fetch it; otherwise save as text."""
    text = text.strip()
    if text.startswith("http://") or text.startswith("https://"):
        out = base_output / "html"
        return convert_html(text, out, vault_dir)
    return convert_raw_text(text, base_output / "txt", vault_dir)
